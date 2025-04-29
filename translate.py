import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from langchain_google_vertexai import VertexAI

class VertexTranslator:
    def __init__(self, llm, target_lang):
        self.llm = llm
        self.target_lang = target_lang
        
    def transform_documents(self, documents):
        translated_docs = []
        for doc in documents:
            # Determine if this is a table
            is_table = doc["type"] == "Table"
            
            # Create appropriate prompt based on element type
            if is_table:
                # Format table for translation
                table_content = self.format_table_for_translation(doc["content"])
                prompt = f"""Translate this table from English to {self.target_lang}. 
                IMPORTANT: Preserve ALL formatting, dollar amounts, numbers, and table structure exactly.
                
                TABLE (tab-separated):
                {table_content}
                
                TRANSLATED TABLE (preserve all tabs and line breaks):"""
            elif doc["type"] == "Heading":
                prompt = f"""Translate this heading from English to {self.target_lang}:
                
                {doc["content"]}
                
                TRANSLATION:"""
            else:
                prompt = f"""Translate this text from English to {self.target_lang}:
                
                {doc["content"]}
                
                TRANSLATION:"""
            
            # Get translation
            translation = self.llm.invoke(prompt)
            
            # Create new document with translation
            translated_doc = doc.copy()
            if is_table:
                translated_doc["content"] = self.parse_translated_table(translation)
            else:
                translated_doc["content"] = translation.strip()
            translated_docs.append(translated_doc)
        
        return translated_docs
    
    def format_table_for_translation(self, table_data):
        """Format a table into a string representation for translation."""
        table_text = ""
        for row in table_data:
            table_text += "\t".join(row) + "\n"
        return table_text
    
    def parse_translated_table(self, translated_text):
        """Parse translated table text back into a nested list structure."""
        rows = translated_text.strip().split('\n')
        table_data = []
        
        for row in rows:
            cells = row.split('\t')
            table_data.append(cells)
            
        return table_data

def extract_document_content(input_path):
    """
    Extract content from a Word document using python-docx
    with special handling for tables.
    """
    doc = Document(input_path)
    document_elements = []
    
    # Extract paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            # Determine if it's a heading
            is_heading = para.style.name.startswith('Heading')
            level = 0
            if is_heading:
                try:
                    # Extract heading level from style name (e.g., 'Heading 1' -> 1)
                    level = int(para.style.name.split()[-1])
                except (ValueError, IndexError):
                    level = 1
            
            document_elements.append({
                'type': 'Heading' if is_heading else 'Paragraph',
                'level': level if is_heading else 0,
                'content': para.text,
                'bold': any(run.bold for run in para.runs),
                'italic': any(run.italic for run in para.runs)
            })
    
    # Extract tables
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                row_data.append(cell.text)
            table_data.append(row_data)
            
        document_elements.append({
            'type': 'Table',
            'content': table_data
        })
    
    return document_elements

def translate_document_with_tables(
    input_path, 
    output_path, 
    project_id, 
    location="us-central1", 
    target_lang="es"
):
    """
    Translate a Word document with tables using python-docx and Vertex AI.
    """
    # Step 1: Extract document content
    print(f"Extracting content from {input_path}...")
    document_elements = extract_document_content(input_path)
    print(f"Extracted {len(document_elements)} elements")
    
    # Step 2: Set up Vertex AI and the translator
    llm = VertexAI(
        model_name="gemini-pro",
        project=project_id,
        location=location,
        temperature=0.1
    )
    
    translator = VertexTranslator(llm, target_lang)
    
    # Step 3: Translate elements
    print("Translating document elements...")
    translated_elements = translator.transform_documents(document_elements)
    print(f"Translated {len(translated_elements)} elements")
    
    # Step 4: Create a new document with translated content
    output_doc = Document()
    
    for element in translated_elements:
        if element['type'] == 'Heading':
            # Add heading with appropriate level
            output_doc.add_heading(element['content'], level=element['level'])
            
        elif element['type'] == 'Paragraph':
            # Add paragraph with formatting
            p = output_doc.add_paragraph(element['content'])
            if element.get('bold'):
                for run in p.runs:
                    run.bold = True
            if element.get('italic'):
                for run in p.runs:
                    run.italic = True
                    
        elif element['type'] == 'Table':
            # Create table with appropriate dimensions
            table_data = element['content']
            if not table_data:
                continue
                
            rows = len(table_data)
            cols = max(len(row) for row in table_data)
            
            table = output_doc.add_table(rows=rows, cols=cols)
            table.style = 'Table Grid'  # Apply a standard table style
            
            # Fill table with translated content
            for i, row in enumerate(table_data):
                for j, cell_text in enumerate(row):
                    if j < cols:
                        table.cell(i, j).text = cell_text
    
    # Save the document
    output_doc.save(output_path)
    print(f"Translated document saved to: {output_path}")
    
    return output_path

# Example usage
if __name__ == "__main__":
    input_file = "business_checking.docx"
    output_file = "business_checking_spanish.docx"
    google_cloud_project_id = "your-google-cloud-project-id"
    
    translate_document_with_tables(
        input_file,
        output_file,
        project_id=google_cloud_project_id
    )
